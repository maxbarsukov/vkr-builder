from docx import Document
from docx.oxml.ns import qn

from vkr import md
from vkr.md import latex_to_omath_element


def _local_names(element):
    names = set()
    for el in element.iter():
        tag = el.tag
        names.add(tag.split("}")[-1] if "}" in tag else tag)
    return names


def test_sum_with_limits_uses_nary():
    om = latex_to_omath_element(r"\sum_{i=1}^{k} 1")
    names = _local_names(om)
    assert "nary" in names
    assert "sSubSup" not in names


def test_combined_sub_and_sup_uses_subsup():
    om = latex_to_omath_element(r"x_i^2")
    names = _local_names(om)
    assert "sSubSup" in names
    assert "sSub" not in names
    assert "sSup" not in names


def test_sub_then_sup_order_is_combined():
    om = latex_to_omath_element(r"x^2_i")
    names = _local_names(om)
    assert "sSubSup" in names


def test_single_subscript_stays_ssub():
    om = latex_to_omath_element(r"x_i")
    names = _local_names(om)
    assert "sSub" in names
    assert "sSubSup" not in names


def _math_ns(tag):
    return "{http://schemas.openxmlformats.org/officeDocument/2006/math}" + tag


def test_nary_operand_is_inside_e_not_empty():
    om = latex_to_omath_element(r"N = \sum_{i=1}^{k} 1")
    nary = om.find(".//" + _math_ns("nary"))
    assert nary is not None
    e = nary.find(_math_ns("e"))
    assert e is not None
    texts = "".join(t.text or "" for t in e.iter(_math_ns("t")))
    assert "1" in texts
    top_texts = [t.text for t in om.findall(_math_ns("r") + "/" + _math_ns("t"))]
    assert "1" not in top_texts


def test_inline_math_is_direct_child_of_paragraph():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("before ")
    md.append_inline_math(p, r"\mu")
    p.add_run(" after")

    children = list(p._element)
    omaths = [c for c in children if c.tag == qn("m:oMath")]
    assert len(omaths) == 1, "inline math should be a direct child of the paragraph"
    for run in p._element.findall(qn("w:r")):
        assert run.find(qn("m:oMath")) is None


def test_nested_fraction_renders_two_fractions():
    om = latex_to_omath_element(r"\frac{1}{1 + \frac{1}{1 + x}}")
    fractions = om.findall(".//" + _math_ns("f"))
    assert len(fractions) == 2


def test_double_sum_nests_two_nary():
    om = latex_to_omath_element(r"\sum_{i=1}^{n} \sum_{j=1}^{n} a_{i}")
    naries = om.findall(".//" + _math_ns("nary"))
    assert len(naries) == 2


def test_bounded_operator_keeps_trailing_terms_outside():
    om = latex_to_omath_element(r"{\int_{0}^{\infty} e^{-x} dx} = 1")
    nary = om.find(".//" + _math_ns("nary"))
    assert nary is not None
    e = nary.find(_math_ns("e"))
    inside = "".join(t.text or "" for t in e.iter(_math_ns("t")))
    assert "1" not in inside
