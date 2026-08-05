import pytest

from vkr.typography import TypographySettings


def test_defaults_match_config():
    t = TypographySettings()
    assert t.body_font_pt == 14.0
    assert t.margin_left_cm == 3.0
    assert t.page_number_from == 6


def test_to_flat_roundtrip():
    t = TypographySettings(body_font_pt=13.0, page_number_from=9)
    back = TypographySettings.from_flat(t.to_flat())
    assert back.body_font_pt == 13.0
    assert back.page_number_from == 9


def test_rejects_small_body_font():
    flat = TypographySettings().to_flat()
    flat["body_font_pt"] = 10
    with pytest.raises(ValueError, match="body_font_pt"):
        TypographySettings.validate(flat)


def test_rejects_invalid_margins():
    flat = TypographySettings().to_flat()
    flat["margin_left_cm"] = 20.0
    with pytest.raises(ValueError, match="margins exceed"):
        TypographySettings.validate(flat)


def test_merge_with_defaults_page_number():
    t = TypographySettings.merge_with_defaults({"body_font_pt": 12.0}, page_number_from=11)
    assert t.body_font_pt == 12.0
    assert t.page_number_from == 11


def test_margins_cm_mapping():
    t = TypographySettings.from_flat({"margins_cm": {"left": 2.5, "right": 2.0}})
    assert t.margin_left_cm == 2.5
    assert t.margin_right_cm == 2.0


def _dash_doc(tmp_path, dashes):
    from vkr import docx_build
    import docx as docxlib

    src = tmp_path / "d.md"
    src.write_text(
        "# ВВЕДЕНИЕ\n\nПервое - дефис, кто-то, научно–технический, "
        "6-7 и 2025 - 2026 гг.\n\n"
        "Таблица {t} - Подпись\n\n| A |\n|---|\n| 1 |\n\nСм. [табл:t].\n",
        encoding="utf-8",
    )
    out = tmp_path / "d.docx"
    docx_build.build(str(src), str(out), include_toc=False, dashes=dashes,
                     pagination_engine="libreoffice")
    doc = docxlib.Document(str(out))
    body = next(p.text for p in doc.paragraphs if "дефис" in p.text)
    caption = next(p.text for p in doc.paragraphs if "Подпись" in p.text)
    return body, caption


def test_itmo_defaults_put_an_en_dash_everywhere(tmp_path):
    body, caption = _dash_doc(tmp_path, {"normalize": True, "captions": "en-dash", "body": "en-dash"})
    assert "Первое – дефис" in body
    assert caption.endswith("– Подпись")


def test_em_dash_is_a_config_switch(tmp_path):
    body, caption = _dash_doc(tmp_path, {"normalize": True, "captions": "em-dash", "body": "em-dash"})
    assert "Первое — дефис" in body
    assert caption.endswith("— Подпись")


def test_captions_and_body_are_independent(tmp_path):
    body, caption = _dash_doc(tmp_path, {"normalize": True, "captions": "en-dash", "body": "em-dash"})
    assert "Первое — дефис" in body
    assert caption.endswith("– Подпись")


def test_normalize_false_keeps_what_the_author_typed(tmp_path):
    body, caption = _dash_doc(tmp_path, {"normalize": False})
    assert "Первое - дефис" in body
    assert caption.endswith("- Подпись")


def test_a_dash_inside_a_word_is_always_a_hyphen(tmp_path):
    body, _ = _dash_doc(tmp_path, {"normalize": True, "body": "em-dash"})
    assert "кто-то" in body
    assert "научно-технический" in body, "en dash typed inside a word is corrected"


def test_a_tight_range_is_always_an_en_dash(tmp_path):
    body, _ = _dash_doc(tmp_path, {"normalize": True, "body": "em-dash"})
    assert "6\u20137" in body


def test_a_spaced_range_keeps_its_spaces(tmp_path):
    body, _ = _dash_doc(tmp_path, {"normalize": True, "body": "em-dash"})
    assert "2025 \u2014 2026" in body, "spacing is never changed, only the glyph"
    assert "2025\u20132026" not in body


def test_a_bad_dash_name_is_a_config_error(tmp_path):
    from vkr import config

    cfg = tmp_path / "c.yaml"
    cfg.write_text(
        "active_profile: p\nprofiles:\n  p: {docx: o.docx, markdown_dir: md,"
        " markdown_files: [a.md]}\nstyle:\n  dashes:\n    body: long\n",
        encoding="utf-8",
    )
    with pytest.raises(config.ConfigError, match="'en-dash' or 'em-dash'"):
        config.load_build_config(cfg)
