from __future__ import annotations

import importlib

from docx.oxml.ns import qn

from vkr import config, docx_style
from vkr.docx import elements as E

build = importlib.import_module("vkr.docx.build")


def _appendix_elements():
    return [
        {"type": "heading", "level": 1, "text": "ПРИЛОЖЕНИЕ В"},
        {"type": "table_caption", "text": "Таблица {elements} – Конструкции"},
        {
            "type": "table",
            "header": ["A", "B"],
            "aligns": ["left", "left"],
            "rows": [[f"r{i}a", f"r{i}b"] for i in range(5)],
        },
        {"type": "table_caption", "text": "Таблица {commands} – Команды"},
        {
            "type": "table",
            "header": ["C"],
            "aligns": ["left"],
            "rows": [["x"], ["y"]],
        },
    ]


class _MockSession:
    def __init__(self, pages: list[list[int]]) -> None:
        self._pages = pages

    def fragment_count(self) -> int:
        return len(self._pages)

    def fragment_row_count(self, fragment_index: int) -> int:
        return len(self._pages[fragment_index])

    def row_page(self, fragment_index: int, data_row_index: int) -> int:
        return self._pages[fragment_index][data_row_index]


def test_table_numbers_use_appendix_labels():
    nums = build._compute_table_numbers(_appendix_elements())
    assert nums == ["В.1", "В.2"]


def test_table_numbers_none_without_caption_key():
    els = [
        {
            "type": "table",
            "header": ["A"],
            "aligns": ["left"],
            "rows": [["1"]],
        }
    ]
    assert build._compute_table_numbers(els) == [None]


def test_enumerate_fragments_from_splits():
    els = [
        {"type": "heading", "level": 1, "text": "1"},
        {
            "type": "table",
            "header": ["H"],
            "aligns": ["left"],
            "rows": [[str(i)] for i in range(6)],
        },
    ]
    frags = build._enumerate_table_fragments(els, {0: (3,)})
    assert len(frags) == 2
    assert frags[0] == build._TableFragment(0, 1, 0, 3, 0)
    assert frags[1] == build._TableFragment(0, 1, 3, 6, 1)


def test_first_row_on_next_page_binary_search():
    pages = [7, 7, 7, 8, 8, 9]
    idx = build._first_row_on_next_page(lambda r: pages[r], len(pages))
    assert idx == 3


def test_first_row_on_next_page_none_when_single_page():
    pages = [5, 5, 5]
    assert build._first_row_on_next_page(lambda r: pages[r], len(pages)) is None


def test_step_marks_single_page_fragment_processed():
    els = _appendix_elements()
    session = _MockSession([[7, 7, 7, 7, 7], [8, 8]])
    processed: set[tuple[int, int]] = set()
    splits, split_made, processed = build._run_table_continuation_step(
        els, {}, processed, session
    )
    assert split_made is False
    assert splits == {}
    assert processed == {(0, 0), (1, 0)}


def test_step_splits_spanning_fragment_once():
    els = _appendix_elements()
    session = _MockSession([[7, 7, 7, 8, 8], [9, 9]])
    processed: set[tuple[int, int]] = set()
    splits, split_made, processed = build._run_table_continuation_step(
        els, {}, processed, session
    )
    assert split_made is True
    assert splits == {0: (3,)}
    assert processed == {(0, 0)}


def test_step_second_pass_splits_remainder():
    els = [
        {"type": "heading", "level": 1, "text": "1"},
        {
            "type": "table",
            "header": ["H"],
            "aligns": ["left"],
            "rows": [[str(i)] for i in range(6)],
        },
    ]
    session = _MockSession([[7, 7, 7], [8, 8, 9]])
    processed = {(0, 0)}
    splits, split_made, processed = build._run_table_continuation_step(
        els, {0: (3,)}, processed, session
    )
    assert split_made is True
    assert splits == {0: (3, 5)}
    assert (0, 3) in processed


def test_step_skips_processed_and_checks_next_in_chapter():
    els = [
        {"type": "heading", "level": 1, "text": "1 Глава"},
        {
            "type": "table",
            "header": ["H"],
            "aligns": ["left"],
            "rows": [[f"a{i}"] for i in range(3)],
        },
        {
            "type": "table",
            "header": ["H"],
            "aligns": ["left"],
            "rows": [[f"b{i}"] for i in range(4)],
        },
    ]
    session = _MockSession([[1, 1, 1], [2, 2, 3, 3]])
    processed = {(0, 0)}
    splits, split_made, processed = build._run_table_continuation_step(
        els, {}, processed, session
    )
    assert split_made is True
    assert splits == {1: (2,)}
    assert processed == {(0, 0), (1, 0)}


def test_step_one_split_per_call_across_chapters():
    els = [
        {"type": "heading", "level": 1, "text": "1"},
        {
            "type": "table",
            "header": ["H"],
            "aligns": ["left"],
            "rows": [[f"a{i}"] for i in range(4)],
        },
        {"type": "heading", "level": 1, "text": "2"},
        {
            "type": "table",
            "header": ["H"],
            "aligns": ["left"],
            "rows": [[f"b{i}"] for i in range(4)],
        },
    ]
    session = _MockSession([[1, 1, 2, 2], [3, 3, 4, 4]])
    splits, split_made, processed = build._run_table_continuation_step(
        els, {}, set(), session
    )
    assert split_made is True
    assert splits == {0: (2,)}
    assert processed == {(0, 0)}


def test_compute_table_splits_full_scan():
    els = _appendix_elements()
    frag_pages = [[7, 7, 7, 8, 8], [9, 9]]
    splits = build._compute_table_splits(els, {}, frag_pages)
    assert splits == {0: (3,)}


def _caption_paragraphs(body):
    out = []
    for p in body.findall(qn("w:p")):
        text = "".join(t.text or "" for t in p.iter(qn("w:t")))
        if text.startswith("Продолжение таблицы"):
            out.append((p, text))
    return out


def test_split_renders_fragments_caption_and_repeated_header():
    docx_style.reset_table_continuation_to_defaults()
    doc = docx_style.create_vkr_document()
    rows = [[f"r{i}a", f"r{i}b"] for i in range(6)]
    E.add_table(
        doc, ["Col A", "Col B"], ["left", "left"], rows,
        split_after=(3,), table_number="В.1",
    )
    body = doc.element.body

    tbls = body.findall(qn("w:tbl"))
    assert len(tbls) == 2

    total_data = 0
    for t in tbls:
        trs = t.findall(qn("w:tr"))
        first_trpr = trs[0].find(qn("w:trPr"))
        assert first_trpr is not None
        assert first_trpr.find(qn("w:tblHeader")) is not None
        total_data += len(trs) - 1
    assert total_data == len(rows)

    captions = _caption_paragraphs(body)
    assert len(captions) == 1
    p, text = captions[0]
    assert text == "Продолжение таблицы В.1"
    pPr = p.find(qn("w:pPr"))
    assert pPr.find(qn("w:pageBreakBefore")) is not None
    jc = pPr.find(qn("w:jc"))
    assert jc is not None and jc.get(qn("w:val")) == "right"


def test_no_split_renders_single_table_without_caption():
    docx_style.reset_table_continuation_to_defaults()
    doc = docx_style.create_vkr_document()
    rows = [[f"r{i}a", f"r{i}b"] for i in range(4)]
    E.add_table(doc, ["A", "B"], ["left", "left"], rows)
    body = doc.element.body
    assert len(body.findall(qn("w:tbl"))) == 1
    assert _caption_paragraphs(body) == []


def test_continuation_alignment_is_configurable():
    docx_style.configure_table_continuation({"align": "left"})
    try:
        doc = docx_style.create_vkr_document()
        rows = [[f"r{i}"] for i in range(4)]
        E.add_table(doc, ["H"], ["left"], rows, split_after=(2,), table_number="1")
        body = doc.element.body
        p, _ = _caption_paragraphs(body)[0]
        jc = p.find(qn("w:pPr")).find(qn("w:jc"))
        assert jc is not None and jc.get(qn("w:val")) == "left"
    finally:
        docx_style.reset_table_continuation_to_defaults()


def test_config_defaults_for_tables():
    cfg = config._parse_table_continuation(None)
    assert cfg.enabled is True
    assert cfg.label == "Продолжение таблицы {n}"
    assert cfg.align == "right"
    assert cfg.to_mapping() == {
        "enabled": True,
        "label": "Продолжение таблицы {n}",
        "align": "right",
    }


def test_config_parses_custom_tables_block():
    cfg = config._parse_table_continuation(
        {
            "continuation": False,
            "continued_label": "Таблица {n} (продолжение)",
            "continuation_align": "center",
        }
    )
    assert cfg.enabled is False
    assert cfg.label == "Таблица {n} (продолжение)"
    assert cfg.align == "center"


def test_config_rejects_bad_alignment():
    try:
        config._parse_table_continuation({"continuation_align": "justify"})
    except config.ConfigError:
        return
    raise AssertionError("expected ConfigError for invalid alignment")


def test_config_rejects_unknown_tables_key():
    try:
        config._parse_table_continuation({"nope": 1})
    except config.ConfigError:
        return
    raise AssertionError("expected ConfigError for unknown key")
