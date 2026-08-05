import os
import sys
from pathlib import Path

import pytest

from vkr import pdf_export


def test_normalize_pdf_engine_accepts_known():
    assert pdf_export.normalize_pdf_engine("word") == "word"
    assert pdf_export.normalize_pdf_engine("LibreOffice") == "libreoffice"


def test_normalize_pdf_engine_rejects_unknown():
    with pytest.raises(ValueError):
        pdf_export.normalize_pdf_engine("acrobat")


def test_default_pdf_path_swaps_extension():
    assert pdf_export.default_pdf_path("a/b/thesis.docx") == Path("a/b/thesis.pdf")


def test_export_pdf_missing_docx(tmp_path):
    with pytest.raises(FileNotFoundError):
        pdf_export.export_pdf(tmp_path / "nope.docx", engine="libreoffice")


class _Conversion:
    returncode = 0
    stdout = "Error: Please verify input parameters..."
    stderr = ""


def _mock_soffice(monkeypatch, writes: bytes | None):
    monkeypatch.setattr(pdf_export, "resolve_libreoffice_path", lambda path: "soffice")
    monkeypatch.setattr(pdf_export.time, "sleep", lambda seconds: None)

    def run(soffice, docx_path, out_dir):
        if writes is not None:
            (out_dir / (docx_path.stem + ".pdf")).write_bytes(writes)
        return _Conversion()

    monkeypatch.setattr(pdf_export, "_run_soffice_convert", run)


def test_stale_output_is_not_mistaken_for_a_conversion(tmp_path, monkeypatch):
    stale = tmp_path / "thesis.pdf"
    stale.write_bytes(b"%PDF-1.7 from an earlier run")
    _mock_soffice(monkeypatch, writes=None)

    with pytest.raises(RuntimeError, match="was not written"):
        pdf_export._export_pdf_libreoffice(tmp_path / "thesis.docx", stale, None)
    assert stale.read_bytes() == b"%PDF-1.7 from an earlier run"


def test_a_rewritten_output_is_accepted(tmp_path, monkeypatch):
    out = tmp_path / "thesis.pdf"
    out.write_bytes(b"%PDF-1.7 from an earlier run")
    _mock_soffice(monkeypatch, writes=b"%PDF-1.7 fresh")

    pdf_export._export_pdf_libreoffice(tmp_path / "thesis.docx", out, None)
    assert out.read_bytes() == b"%PDF-1.7 fresh"


@pytest.mark.skipif(
    sys.platform == "win32", reason="POSIX permission bits behave differently"
)
def test_unwritable_output_fails_without_retrying(tmp_path, monkeypatch):
    locked = tmp_path / "thesis.pdf"
    locked.write_bytes(b"%PDF-1.7 read-only")
    locked.chmod(0o444)
    calls = []
    _mock_soffice(monkeypatch, writes=None)
    monkeypatch.setattr(
        pdf_export,
        "_run_soffice_convert",
        lambda *args: (calls.append(1), _Conversion())[1],
    )

    with pytest.raises(RuntimeError, match="cannot be written to"):
        pdf_export._export_pdf_libreoffice(tmp_path / "thesis.docx", locked, None)
    assert calls == []


def _libreoffice_available():
    try:
        from vkr.pagination import resolve_libreoffice_path

        resolve_libreoffice_path(None)
        return True
    except Exception:
        return False


@pytest.mark.skipif(
    not (os.environ.get("VKR_TEST_PDF") and _libreoffice_available()),
    reason="set VKR_TEST_PDF=1 and install LibreOffice to run the real conversion",
)
def test_libreoffice_conversion_creates_pdf(tmp_path):
    from docx import Document

    src = tmp_path / "small.docx"
    doc = Document()
    doc.add_paragraph("Hello PDF")
    doc.save(str(src))

    out = pdf_export.export_pdf(src, engine="libreoffice")
    assert out.is_file()
    assert out.stat().st_size > 0


def test_uno_interpreter_falls_back_to_a_system_python(tmp_path, monkeypatch):
    from vkr import pagination

    program = tmp_path / "program"
    program.mkdir()
    soffice = program / "soffice"
    soffice.write_text("", encoding="utf-8")

    monkeypatch.setattr(pagination.shutil, "which", lambda name: None)
    monkeypatch.setattr(pagination.glob, "glob", lambda pattern: ["/usr/bin/python3.13"])
    monkeypatch.setattr(pagination.os.path, "isfile", lambda path: True)
    monkeypatch.setattr(
        pagination, "_can_import_uno", lambda path: path == "/usr/bin/python3.13"
    )
    pagination._UNO_PROBE_CACHE.clear()

    assert pagination._lo_python_path(str(soffice)) == "/usr/bin/python3.13"


def test_uno_interpreter_prefers_the_one_next_to_soffice(tmp_path, monkeypatch):
    from vkr import pagination

    program = tmp_path / "program"
    program.mkdir()
    (program / "soffice").write_text("", encoding="utf-8")
    bundled = program / "python"
    bundled.write_text("", encoding="utf-8")

    monkeypatch.setattr(
        pagination, "_can_import_uno", lambda path: pytest.fail("must not probe")
    )
    assert pagination._lo_python_path(str(program / "soffice")) == str(bundled)


def test_missing_uno_bridge_says_what_to_install(tmp_path, monkeypatch):
    from vkr import pagination

    program = tmp_path / "program"
    program.mkdir()
    soffice = program / "soffice"
    soffice.write_text("", encoding="utf-8")

    monkeypatch.setattr(pagination.shutil, "which", lambda name: None)
    monkeypatch.setattr(pagination.glob, "glob", lambda pattern: [])
    monkeypatch.setattr(pagination, "_can_import_uno", lambda path: False)
    pagination._UNO_PROBE_CACHE.clear()

    with pytest.raises(FileNotFoundError) as exc:
        pagination._lo_python_path(str(soffice))
    assert "python3-uno" in str(exc.value)


def test_word_export_asks_for_a_navigable_tree(tmp_path, monkeypatch):
    import contextlib

    asked = {}

    class _Document:
        def ExportAsFixedFormat(self, **kwargs):
            asked.update(kwargs)
            Path(kwargs["OutputFileName"]).write_bytes(b"%PDF-1.7\n")

    @contextlib.contextmanager
    def _application(purpose):
        yield object()

    @contextlib.contextmanager
    def _document(word, path, purpose):
        yield _Document()

    monkeypatch.setattr(pdf_export, "word_application", _application)
    monkeypatch.setattr(pdf_export, "open_word_document", _document)

    out = tmp_path / "thesis.pdf"
    pdf_export._export_pdf_word(tmp_path / "thesis.docx", out)

    assert asked["CreateBookmarks"] == pdf_export._WD_EXPORT_CREATE_HEADING_BOOKMARKS
    assert asked["CreateBookmarks"] != 2
    assert asked["DocStructureTags"] is True
    assert out.is_file()
