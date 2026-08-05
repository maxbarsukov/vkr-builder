import zipfile

from vkr import docx_build
from vkr import md
from tests.docx_inspect import read_docx


def _build(tmp_path, md_path):
    elements = md.parse_md(str(md_path))
    out = tmp_path / "out.docx"
    docx_build._build_pass(str(out), elements, None, assets_root=None)
    return out


def test_build_produces_readable_docx(tmp_path, sample_md_path):
    out = _build(tmp_path, sample_md_path)
    assert out.is_file()
    view = read_docx(out)
    assert view.paragraphs, "document should contain paragraphs"


def test_build_renders_table(tmp_path, sample_md_path):
    out = _build(tmp_path, sample_md_path)
    view = read_docx(out)
    assert view.table_count >= 2


def test_table_header_repeats_on_page_break(tmp_path, sample_md_path):
    out = _build(tmp_path, sample_md_path)
    view = read_docx(out)
    assert "<w:tblHeader" in view.document_xml
    assert "<w:cantSplit" in view.document_xml


def test_build_renders_math(tmp_path, sample_md_path):
    out = _build(tmp_path, sample_md_path)
    view = read_docx(out)
    assert view.math_count >= 2


def test_build_keeps_headings(tmp_path, sample_md_path):
    out = _build(tmp_path, sample_md_path)
    view = read_docx(out)
    assert view.has("Test chapter")
    assert view.has("Lists")


def test_build_keeps_list_items(tmp_path, sample_md_path):
    out = _build(tmp_path, sample_md_path)
    view = read_docx(out)
    assert view.has("first numbered item")
    assert view.has("first bullet item")


NESTED_LIST_MD = """\
# Chapter

The tool supports nested enumerations:

- top level item one:
  1) nested numbered one;
  2) nested numbered two;
- top level item two.
"""


def test_nested_list_uses_second_level(tmp_path):
    import zipfile

    src = tmp_path / "nested.md"
    src.write_text(NESTED_LIST_MD, encoding="utf-8")
    elements = md.parse_md(str(src))

    levels = [e.get("level") for e in elements if e["type"] == "list_item"]
    assert levels == [0, 1, 1, 0]

    out = tmp_path / "out.docx"
    docx_build._build_pass(str(out), elements, None, assets_root=None)

    view = read_docx(out)
    assert view.document_xml.count('<w:ilvl w:val="1"') == 2

    with zipfile.ZipFile(str(out)) as zf:
        numbering = zf.read("word/numbering.xml").decode("utf-8")
    assert "multiLevel" in numbering


def test_metadata_applied_to_core_properties(tmp_path, sample_md_path):
    from docx import Document

    elements = md.parse_md(str(sample_md_path))
    out = tmp_path / "meta.docx"
    docx_build._DOC_METADATA = {
        "title": "My Title",
        "author": "Jane Doe",
        "created": "2026-01-15",
    }
    try:
        docx_build._build_pass(str(out), elements, None, assets_root=None)
    finally:
        docx_build._DOC_METADATA = None

    cp = Document(str(out)).core_properties
    assert cp.title == "My Title"
    assert cp.author == "Jane Doe"
    assert cp.created.year == 2026 and cp.created.month == 1


def _zip_dates(path):
    with zipfile.ZipFile(str(path)) as zf:
        return {info.filename: info.date_time for info in zf.infolist()}


def _build_without_an_engine(out, sample_md_path, metadata=None):
    md = out.with_suffix(".md")
    md.write_text("Body paragraph.\n", encoding="utf-8")
    return docx_build.build(
        str(md), str(out), include_toc=False,
        pagination_engine="libreoffice", metadata=metadata,
    )


def test_container_dates_follow_the_configured_modification_date(tmp_path, sample_md_path):
    out = tmp_path / "dated.docx"
    _build_without_an_engine(
        out, sample_md_path, metadata={"modified": "2026-06-23 18:45:00"}
    )

    dates = set(_zip_dates(out).values())
    assert dates == {(2026, 6, 23, 18, 45, 0)}


def test_container_dates_fall_back_to_the_neutral_one(tmp_path, sample_md_path):
    from vkr.config import NEUTRAL_TIMESTAMP

    out = tmp_path / "neutral.docx"
    _build_without_an_engine(out, sample_md_path, metadata={"title": "No dates here"})

    neutral = NEUTRAL_TIMESTAMP
    assert set(_zip_dates(out).values()) == {
        (neutral.year, neutral.month, neutral.day,
         neutral.hour, neutral.minute, neutral.second)
    }


def test_two_builds_of_the_same_markdown_are_the_same_file(tmp_path, sample_md_path):
    metadata = {"title": "Reproducible", "modified": "2026-06-23"}
    outputs = []
    for name in ("first.docx", "second.docx"):
        out = tmp_path / name
        _build_without_an_engine(out, sample_md_path, metadata=metadata)
        outputs.append(out.read_bytes())

    assert outputs[0] == outputs[1]


def test_repacking_keeps_the_parts_and_their_order(tmp_path, sample_md_path):
    from docx import Document

    out = tmp_path / "order.docx"
    _build_without_an_engine(out, sample_md_path, metadata={"modified": "2026-06-23"})

    with zipfile.ZipFile(str(out)) as zf:
        names = zf.namelist()
        assert zf.testzip() is None
    assert names[0] == "[Content_Types].xml"
    assert "word/document.xml" in names
    assert Document(str(out)).paragraphs


def test_a_date_before_1980_is_pushed_into_what_a_zip_can_hold(tmp_path, sample_md_path):
    import datetime as dt

    from vkr.docx.document import normalise_package_timestamps

    out = tmp_path / "old.docx"
    _build_without_an_engine(out, sample_md_path)
    normalise_package_timestamps(out, dt.datetime(1970, 5, 6, 7, 8, 9))

    assert set(_zip_dates(out).values()) == {(1980, 1, 1, 0, 0, 0)}


def test_metadata_cleanup_when_unset(tmp_path, sample_md_path):
    from docx import Document

    elements = md.parse_md(str(sample_md_path))
    out = tmp_path / "clean.docx"
    docx_build._DOC_METADATA = None
    docx_build._build_pass(str(out), elements, None, assets_root=None)

    cp = Document(str(out)).core_properties
    assert not cp.author
    assert not cp.title
