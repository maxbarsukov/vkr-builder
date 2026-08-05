from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document

from .docx_inspect import read_docx
from .logging_setup import get_logger
from .word_com import word_document_session

log = get_logger("diagnostics")

_FIGURE_CAPTION_RE = re.compile(r"^\s*Рисунок\s+\S", re.IGNORECASE)
_TABLE_CAPTION_RE = re.compile(r"^\s*Таблица\s+\S", re.IGNORECASE)
_HEADING_STYLE_PREFIX = "Heading"


@dataclass(frozen=True)
class DiagnosticIssue:
    severity: str
    message: str
    location: str = ""
    rule: str = ""


def _figure_caption_count(doc: Document) -> int:
    return sum(1 for p in doc.paragraphs if _FIGURE_CAPTION_RE.match(p.text or ""))


def _table_caption_count(doc: Document) -> int:
    return sum(1 for p in doc.paragraphs if _TABLE_CAPTION_RE.match(p.text or ""))


def diagnose_docx(
    path: Path | str,
    *,
    max_paragraph_chars: int = 2000,
) -> list[DiagnosticIssue]:
    path = Path(path)
    issues: list[DiagnosticIssue] = []
    view = read_docx(path)
    doc = Document(str(path))

    figures = view.drawing_count
    fig_caps = _figure_caption_count(doc)
    if figures > fig_caps:
        issues.append(DiagnosticIssue(
            "warning",
            f"figures without captions: {figures - fig_caps} of {figures} "
            f"images have no caption",
            rule="figure-captions",
        ))
    if fig_caps > figures:
        issues.append(DiagnosticIssue(
            "warning",
            f"figure captions without an image: {fig_caps - figures} extra",
            rule="figure-captions",
        ))

    tables = view.table_count
    tbl_caps = _table_caption_count(doc)
    if tbl_caps > 0 and tables < tbl_caps:
        issues.append(DiagnosticIssue(
            "warning",
            f"more table captions ({tbl_caps}) than tables ({tables})",
            rule="table-captions",
        ))

    for idx, para in enumerate(doc.paragraphs, start=1):
        text = (para.text or "").strip()
        style = para.style.name if para.style is not None else ""
        if style.startswith(_HEADING_STYLE_PREFIX) and not text:
            issues.append(DiagnosticIssue(
                "warning",
                f"empty heading (style {style!r})",
                f"paragraph {idx}",
                "empty-heading",
            ))
        if len(para.text or "") > max_paragraph_chars:
            issues.append(DiagnosticIssue(
                "warning",
                f"very long paragraph: {len(para.text)} characters "
                f"(limit {max_paragraph_chars})",
                f"paragraph {idx}",
                "long-paragraph",
            ))

    prev_was_heading = False
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        style = para.style.name if para.style is not None else ""
        is_heading = style.startswith(_HEADING_STYLE_PREFIX)
        if prev_was_heading and is_heading:
            issues.append(DiagnosticIssue(
                "warning",
                "section with no body text (consecutive headings)",
                rule="empty-section",
            ))
            break
        prev_was_heading = is_heading and bool(text)

    return issues


def check_orphan_widow(
    path: Path | str,
    *,
    engine: str = "word",
    libreoffice_path: str | None = None,
) -> list[DiagnosticIssue]:
    path = Path(path)
    try:
        if engine == "word":
            return _orphan_widow_word(path)
        if engine == "libreoffice":
            return _orphan_widow_libreoffice(path, libreoffice_path)
    except Exception as exc:
        log.info("orphan/widow check skipped: %s", exc)
    return []


def _orphan_widow_word(path: Path) -> list[DiagnosticIssue]:
    import sys

    if sys.platform != "win32":
        return []

    issues: list[DiagnosticIssue] = []
    log.debug("Orphan/widow check (Word): %s", path.resolve())
    with word_document_session(path, purpose="orphan-widow") as doc:
        for i in range(1, doc.Paragraphs.Count + 1):
            para = doc.Paragraphs(i)
            if para.Range.Words.Count <= 2:
                page = para.Range.Information(3)
                lines = para.Range.ComputeStatistics(2)
                if lines == 1:
                    log.debug(
                        "Orphan/widow candidate: paragraph %d page %s",
                        i, page,
                    )
                    issues.append(DiagnosticIssue(
                        "warning",
                        f"possible orphan/widow line (page {page})",
                        f"paragraph {i}",
                        "orphan-widow",
                    ))
    log.debug("Orphan/widow check (Word): %d issue(s)", len(issues))
    return issues


def _orphan_widow_libreoffice(path: Path, libreoffice_path: str | None) -> list[DiagnosticIssue]:
    return []


def run_diagnostics(
    path: Path | str,
    *,
    max_paragraph_chars: int = 2000,
    check_orphans: bool = True,
    pagination_engine: str = "word",
    libreoffice_path: str | None = None,
) -> list[DiagnosticIssue]:
    issues = diagnose_docx(path, max_paragraph_chars=max_paragraph_chars)
    if check_orphans:
        issues.extend(
            check_orphan_widow(
                path,
                engine=pagination_engine,
                libreoffice_path=libreoffice_path,
            )
        )
    return issues
