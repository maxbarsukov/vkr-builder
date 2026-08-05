from __future__ import annotations

import zipfile
from dataclasses import dataclass, field

from docx import Document

_MATH_TAG = "<m:oMath"
_DRAWING_TAG = "<w:drawing"


@dataclass
class DocxView:
    paragraphs: list[str]
    table_count: int
    media_files: list[str]
    math_count: int
    drawing_count: int
    document_xml: str = field(repr=False)

    @property
    def text(self) -> str:
        return "\n".join(self.paragraphs)

    def has(self, needle: str) -> bool:
        return needle in self.text


def read_docx(path) -> DocxView:
    path = str(path)
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs]
    table_count = len(doc.tables)
    with zipfile.ZipFile(path) as zf:
        media = sorted(n for n in zf.namelist() if n.startswith("word/media/"))
        document_xml = zf.read("word/document.xml").decode("utf-8")
    return DocxView(
        paragraphs=paragraphs,
        table_count=table_count,
        media_files=media,
        math_count=document_xml.count(_MATH_TAG),
        drawing_count=document_xml.count(_DRAWING_TAG),
        document_xml=document_xml,
    )
